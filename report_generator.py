"""Report generator — produces a .docx report from checklist items."""
from __future__ import annotations

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from models import ChecklistItem

logger = logging.getLogger(__name__)

# Status display labels
_STATUS_LABELS = {
    "sufficient": "Достаточно",
    "unclear": "Неясно",
    "insufficient": "Недостаточно",
    "confirmed": "Подтверждено",
    "unconfirmed": "Не подтверждено",
    "unknown": "Не оценено",
}


def _status_label(status: str) -> str:
    return _STATUS_LABELS.get(status, status)


def _score_bar(score: Optional[float]) -> str:
    if score is None:
        return "—"
    blocks = int(score * 10)
    return "█" * blocks + "░" * (10 - blocks) + f" {score:.2f}"


def generate_report(
    items: List[ChecklistItem],
    output_path: str,
    act_name: str = "",
    pipeline_stats: Optional[Dict] = None,
) -> str:
    """Generate a .docx report from checklist items.

    Returns the path to the generated file.
    """
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        _generate_docx(items, output_path, act_name, pipeline_stats)
    except ImportError:
        logger.warning("python-docx not available — generating markdown report instead")
        md_path = str(output_path).replace(".docx", ".md")
        _generate_markdown(items, md_path, act_name, pipeline_stats)
        return md_path

    return output_path


def _generate_docx(
    items: List[ChecklistItem],
    output_path: str,
    act_name: str,
    pipeline_stats: Optional[Dict],
) -> None:
    import docx
    from docx.shared import Pt, RGBColor

    doc = docx.Document()

    # Title
    title = doc.add_heading("Чеклист нарушений", level=0)
    title.alignment = 1  # center

    if act_name:
        doc.add_paragraph(f"Акт: {act_name}")
    doc.add_paragraph(f"Дата отчёта: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Нарушений обработано: {len(items)}")

    doc.add_heading("Содержание нарушений", level=1)

    ungrounded_items: List[ChecklistItem] = []

    for i, item in enumerate(items, start=1):
        doc.add_heading(f"Нарушение {i}", level=2)

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_row(label: str, value: str) -> None:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value or "—"

        add_row("Источник", item.source_document)
        add_row("Страница", str(item.page))
        add_row("Раздел", item.section)
        add_row("Субъект", item.subject)
        add_row("Норма", item.law_ref)
        add_row("Формулировка", item.raw_text)
        add_row("Полный контекст", item.description)
        add_row("Улучшенная формулировка", item.improved_formulation[:200] if item.improved_formulation else "—")
        add_row("Правовая квалификация", item.legal_qualification or "—")
        add_row("Квалификация подтверждена", "Да" if item.legal_qualification_grounded else "Нет ⚠")
        add_row("Рекомендация", item.recommendation[:200] if item.recommendation else "—")
        add_row("", "")
        add_row("Оценка доказательности", f"{_score_bar(item.evidence_score)} ({_status_label(item.evidence_status)})")
        add_row("Правовая корректность", f"{_score_bar(item.legal_score)} ({_status_label(item.legal_status)})")
        add_row("Исполнимость", f"{_score_bar(item.actionability_score)} ({_status_label(item.actionability_status)})")
        add_row("Уверенность", _score_bar(item.confidence_score))
        add_row("Возможно не нарушение", "Да ⚠" if item.possibly_not_a_violation else "Нет")

        if item.trace:
            add_row("Использованные чанки", ", ".join(item.trace.used_chunk_ids[:5]) or "—")

        doc.add_paragraph()

        if not item.legal_qualification_grounded:
            ungrounded_items.append(item)

    # Ungrounded norms section
    if ungrounded_items:
        doc.add_page_break()
        doc.add_heading("Требуют ручной проверки — неподтверждённые нормы", level=1)
        doc.add_paragraph(
            "Следующие нарушения содержат правовые квалификации, которые не были "
            "подтверждены в нормативной базе. Необходима ручная проверка."
        )
        for item in ungrounded_items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"Нарушение {item.violation_id[:8]}...: ").bold = True
            p.add_run(item.legal_qualification or item.law_ref)

    # Pipeline stats section
    if pipeline_stats:
        doc.add_page_break()
        doc.add_heading("Статистика обработки", level=1)
        stats_table = doc.add_table(rows=0, cols=2)
        stats_table.style = "Table Grid"
        for key, value in pipeline_stats.items():
            row = stats_table.add_row()
            row.cells[0].text = key.replace("_", " ").title()
            row.cells[1].text = str(value)

    doc.save(output_path)
    logger.info("Report saved to %s", output_path)


def _generate_markdown(
    items: List[ChecklistItem],
    output_path: str,
    act_name: str,
    pipeline_stats: Optional[Dict],
) -> None:
    lines = [
        "# Чеклист нарушений",
        "",
        f"**Акт:** {act_name}" if act_name else "",
        f"**Дата отчёта:** {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        f"**Нарушений обработано:** {len(items)}",
        "",
        "---",
        "",
    ]

    ungrounded: List[ChecklistItem] = []

    for i, item in enumerate(items, start=1):
        lines.extend([
            f"## Нарушение {i}",
            "",
            f"| Поле | Значение |",
            f"|---|---|",
            f"| Источник | {item.source_document} |",
            f"| Страница | {item.page} |",
            f"| Субъект | {item.subject} |",
            f"| Норма | {item.law_ref} |",
            f"| Описание | {item.description[:150]} |",
            f"| Улучшенная формулировка | {item.improved_formulation[:150] if item.improved_formulation else '—'} |",
            f"| Квалификация | {item.legal_qualification or '—'} |",
            f"| Квалификация подтверждена | {'✅ Да' if item.legal_qualification_grounded else '⚠️ Нет'} |",
            f"| Рекомендация | {item.recommendation[:150] if item.recommendation else '—'} |",
            f"| Доказательность | {item.evidence_score:.2f if item.evidence_score is not None else '—'} ({_status_label(item.evidence_status)}) |",
            f"| Правовая корректность | {item.legal_score:.2f if item.legal_score is not None else '—'} ({_status_label(item.legal_status)}) |",
            f"| Исполнимость | {item.actionability_score:.2f if item.actionability_score is not None else '—'} ({_status_label(item.actionability_status)}) |",
            f"| Уверенность | {item.confidence_score:.2f} |",
            f"| Возможно не нарушение | {'⚠️ Да' if item.possibly_not_a_violation else 'Нет'} |",
            "",
        ])

        if not item.legal_qualification_grounded:
            ungrounded.append(item)

    if ungrounded:
        lines.extend([
            "---",
            "",
            "## ⚠️ Требуют ручной проверки — неподтверждённые нормы",
            "",
            "Следующие нарушения содержат правовые квалификации, не подтверждённые в нормативной базе:",
            "",
        ])
        for item in ungrounded:
            lines.append(f"- **{item.violation_id[:8]}...**: {item.legal_qualification or item.law_ref}")
        lines.append("")

    if pipeline_stats:
        lines.extend([
            "---",
            "",
            "## Статистика обработки",
            "",
            "| Параметр | Значение |",
            "|---|---|",
        ])
        for key, value in pipeline_stats.items():
            lines.append(f"| {key.replace('_', ' ').title()} | {value} |")

    Path(output_path).write_text("\n".join(lines), encoding="utf-8")
    logger.info("Markdown report saved to %s", output_path)


def generate_report_bytes(
    items: List[ChecklistItem],
    act_name: str = "",
    pipeline_stats: Optional[Dict] = None,
) -> bytes:
    """Generate a .docx report and return it as bytes (for Streamlit download)."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name

    generate_report(items, path, act_name=act_name, pipeline_stats=pipeline_stats)
    return Path(path).read_bytes()
