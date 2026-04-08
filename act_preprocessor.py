"""Act preprocessor — structured violation extraction from DOCX tables and text.

Three-path strategy:
1. DOCX table extraction — reads violations directly from the structured table
2. Text segmentation — regex on resolutive zone only (PDF / plain-text fallback)
3. LLM fallback — full document passed to LLM when no structure is recognised
"""
from __future__ import annotations

import logging
import re
import uuid
from dataclasses import dataclass
from typing import Dict, List, Optional

from models import Violation

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Compiled patterns
# ---------------------------------------------------------------------------

# Resolutive section header markers
_RESOLUTIVE_RE = re.compile(
    r"(?im)^\s*("
    r"резолютивная\s+часть"
    r"|перечень\s+нарушений"
    r"|выявленные\s+нарушения"
    r"|нарушения\s+законодательства"
    r"|нарушения"
    r")[:\s]*$"
)

# Descriptive section header markers
_DESCRIPTIVE_RE = re.compile(
    r"(?im)^\s*("
    r"описательная\s+часть"
    r"|мотивировочная\s+часть"
    r"|установлено"
    r"|результаты\s+проверки"
    r")[:\s]*$"
)

# Numbered item in resolutive text: "1. text..." or "1) text..."
_NUMBERED_ITEM_RE = re.compile(
    r"(?m)^(\d+)[.)]\s+(.{20,}?)(?=\n\d+[.)]\s|\Z)",
    re.DOTALL,
)

# Law reference extraction from raw text
_LAW_REF_RE = re.compile(
    r"(?:Федеральный\s+закон|ФЗ|Постановление|приказ|ГОСТ|СП|СНиП|КоАП"
    r"|ТК\s+РФ|ГК\s+РФ|ГрК\s+РФ|УК\s+РФ)"
    r"[\s\-—]*(?:№\s*)?[\d\-]+[^\s,;.]{0,40}",
    re.IGNORECASE | re.UNICODE,
)

# Subject / organisation pattern
_SUBJECT_RE = re.compile(
    r"(?:ООО|АО|ПАО|ГУП|МУП|ИП|ФГУП|ФГБНУ|ФГБОУ|ФГАОУ)[^,;\n]{3,60}",
    re.UNICODE | re.IGNORECASE,
)

# Distinctive tokens within a law reference (numbers, codes, article refs)
_LAW_TOKEN_RE = re.compile(
    r"[А-ЯЁA-Za-z]{2,}-\d+"   # e.g. ФЗ-69, ФЗ-116, ФЗ-402 (letters-hyphen-number)
    r"|\d+[-/]\w+"             # e.g. 402-ФЗ, 28/в7 (number-hyphen-letters)
    r"|\d+[а-яёА-ЯЁa-zA-Z]+"  # e.g. 157н
    r"|\d+\.\d+"               # e.g. 55.24
    r"|[А-ЯЁA-Z]{2,}\s*РФ"    # e.g. ТК РФ, ГК РФ
    r"|\b\d{3,}\b",            # standalone numbers ≥ 3 digits
    re.UNICODE,
)

# Expected column headers in the violations table (normalised)
_VIOLATIONS_TABLE_HEADERS = frozenset({
    "вид нарушения",
    "нарушенные нпа",
    "формулировка нарушения",
})

# Roman-numeral theme row prefix pattern
_ROMAN_PREFIX_RE = re.compile(r"^(I{1,3}|IV|V?I{0,3}|IX|X)\b", re.UNICODE)

# Anchor pattern: paragraph beginning with "В нарушение..."
_VIOLATION_ANCHOR_RE = re.compile(r"(?i)^\s*в\s+нарушени[еи]")


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class ViolationWithContext:
    """Intermediate object: one violation statement paired with its full context.

    statement         — "Вид нарушения" column (short category label)
    table_formulation — "Формулировка нарушения" column (brief summary from table)
    context           — paragraph block from the descriptive section
                        (facts + formal violation statement + responsible persons)
    """

    violation_id: str
    statement: str
    table_formulation: str
    context: str
    section: str
    page: int
    source_document: str
    law_ref: str
    subject: str

    def __post_init__(self) -> None:
        if not self.violation_id:
            self.violation_id = str(uuid.uuid4())


@dataclass
class _ViolationRow:
    """Internal: one parsed violation row extracted from the table."""

    num: str
    violation_type: str
    law_ref: str
    formulation: str
    current_section: str


# ---------------------------------------------------------------------------
# DOCX path — helper functions
# ---------------------------------------------------------------------------


def _find_violations_table(doc):
    """Return the violations table from a python-docx Document, or None.

    Identified by a first row containing all three expected column names.
    Not tied to any fixed table index so it works across different act versions.
    """
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 4:
            continue
        header_texts = {c.text.strip().lower() for c in table.rows[0].cells}
        if _VIOLATIONS_TABLE_HEADERS.issubset(header_texts):
            logger.info(
                '{"event": "PREPROCESSOR_TABLE_FOUND", "rows": %d, "cols": %d}',
                len(table.rows), len(table.columns),
            )
            return table
    return None


def _get_descriptive_paragraphs(doc) -> List[str]:
    """Return paragraph texts from the descriptive section only.

    Collects all non-empty paragraphs up to (but not including) the first
    paragraph matching the resolutive section marker.
    """
    _RESOLUTIVE_PARA_RE = re.compile(r"(?i)резолютивная\s+часть")
    result = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if _RESOLUTIVE_PARA_RE.search(text):
            break
        result.append(text)
    return result


def _parse_table_rows(table) -> List[_ViolationRow]:
    """Parse the violations table into a list of _ViolationRow objects.

    Skips the header row (row 0), merged theme rows, and status rows.
    Tracks the current theme section for attribution to each violation.
    """
    current_section = ""
    rows: List[_ViolationRow] = []

    for ri, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue

        if ri == 0:
            continue  # header row

        # Merged row: all non-empty cells contain the same text
        non_empty = [c for c in cells if c]
        if non_empty and len(set(non_empty)) == 1:
            text = non_empty[0]
            # Update section only for Roman-numeral theme rows
            if _ROMAN_PREFIX_RE.match(text):
                current_section = text
            continue

        # Violation row: first cell is a bare number (e.g. "1", "2.", "3")
        if re.match(r"^\d+\.?$", cells[0]):
            rows.append(_ViolationRow(
                num=cells[0],
                violation_type=cells[1] if len(cells) > 1 else "",
                law_ref=cells[2] if len(cells) > 2 else "",
                formulation=cells[3] if len(cells) > 3 else "",
                current_section=current_section,
            ))

    return rows


def _normalize_for_match(text: str) -> str:
    """Lowercase and collapse all whitespace including non-breaking spaces."""
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).lower().strip()


def _extract_law_tokens(law_ref: str) -> List[str]:
    """Extract distinctive tokens from a law reference string for matching."""
    normalized = law_ref.replace("\xa0", " ")
    tokens = _LAW_TOKEN_RE.findall(normalized)
    seen: set = set()
    result: List[str] = []
    for t in tokens:
        tn = _normalize_for_match(t)
        if tn and tn not in seen:
            seen.add(tn)
            result.append(tn)
    return result


def _find_context_in_descriptive(
    law_ref: str,
    paragraphs: List[str],
    table_formulation: str,
    before: int = 5,
    after: int = 3,
) -> str:
    """Find the full contextual paragraph block for a violation.

    Strategy:
    1. Extract distinctive tokens from law_ref.
    2. Find the "В нарушение..." paragraph with the highest token overlap.
    3. Return a window of [before] paragraphs before the anchor,
       the anchor itself, and [after] paragraphs after it.
    4. Fall back to table_formulation if no anchor is found.
    """
    tokens = _extract_law_tokens(law_ref)
    if not tokens:
        return table_formulation

    best_idx: Optional[int] = None
    best_score = 0

    for i, para in enumerate(paragraphs):
        if not _VIOLATION_ANCHOR_RE.match(para):
            continue
        para_norm = _normalize_for_match(para)
        score = sum(1 for t in tokens if t in para_norm)
        if score > best_score:
            best_score = score
            best_idx = i

    if best_idx is None or best_score == 0:
        logger.debug(
            '{"event": "PREPROCESSOR_CONTEXT_MISS", "law_ref": "%s"}',
            law_ref[:80],
        )
        return table_formulation

    start = max(0, best_idx - before)
    end = min(len(paragraphs), best_idx + after + 1)
    block = "\n".join(p for p in paragraphs[start:end] if p)

    logger.debug(
        '{"event": "PREPROCESSOR_CONTEXT_HIT", "anchor_idx": %d, "window": [%d, %d]}',
        best_idx, start, end,
    )
    return block


def _extract_subject(text: str) -> str:
    m = _SUBJECT_RE.search(text)
    return m.group(0).strip() if m else ""


def _extract_from_docx(file_path: str, source_doc: str) -> List[ViolationWithContext]:
    """DOCX path: extract violations directly from the structured violations table."""
    try:
        import docx as python_docx
    except ImportError:
        logger.warning(
            '{"event": "PREPROCESSOR_DOCX_UNAVAILABLE", "source": "%s"}',
            source_doc,
        )
        return []

    try:
        doc = python_docx.Document(file_path)
    except Exception as exc:
        logger.error(
            '{"event": "PREPROCESSOR_DOCX_OPEN_ERROR", "error": "%s", "source": "%s"}',
            str(exc), source_doc,
        )
        return []

    table = _find_violations_table(doc)
    if table is None:
        logger.info(
            '{"event": "PREPROCESSOR_NO_TABLE", "source": "%s"}',
            source_doc,
        )
        return []

    desc_paragraphs = _get_descriptive_paragraphs(doc)
    violation_rows = _parse_table_rows(table)

    items: List[ViolationWithContext] = []
    for row in violation_rows:
        ctx = _find_context_in_descriptive(
            row.law_ref, desc_paragraphs, row.formulation
        )
        subject = _extract_subject(ctx) or _extract_subject(row.formulation)
        items.append(ViolationWithContext(
            violation_id=str(uuid.uuid4()),
            statement=row.violation_type,
            table_formulation=row.formulation,
            context=ctx,
            section=row.current_section,
            page=1,
            source_document=source_doc,
            law_ref=row.law_ref,
            subject=subject,
        ))

    return items


# ---------------------------------------------------------------------------
# Text path — helper functions (PDF / plain-text fallback)
# ---------------------------------------------------------------------------


def _segment_act(text: str) -> Dict[str, str]:
    """Split act text into named zones: 'descriptive' and 'resolutive'.

    Returns a dict; values are empty strings when a zone is not found.
    """
    descriptive_match = _DESCRIPTIVE_RE.search(text)
    resolutive_match = _RESOLUTIVE_RE.search(text)

    descriptive_text = ""
    resolutive_text = ""

    if descriptive_match and resolutive_match:
        d_end = descriptive_match.end()
        r_start = resolutive_match.start()
        if d_end < r_start:
            descriptive_text = text[d_end:r_start].strip()
            resolutive_text = text[resolutive_match.end():].strip()
        else:
            # Resolutive marker appears before the descriptive one — unusual layout
            resolutive_text = text[resolutive_match.end():].strip()
            descriptive_text = text[:resolutive_match.start()].strip()
    elif resolutive_match:
        resolutive_text = text[resolutive_match.end():].strip()
        descriptive_text = text[:resolutive_match.start()].strip()
    elif descriptive_match:
        descriptive_text = text[descriptive_match.end():].strip()

    return {"descriptive": descriptive_text, "resolutive": resolutive_text}


def _extract_from_text(raw_text: str, source_doc: str) -> List[ViolationWithContext]:
    """Text path: apply regex only to the resolutive zone of the plain text."""
    segments = _segment_act(raw_text)
    resolutive = segments.get("resolutive", "")
    descriptive = segments.get("descriptive", "")

    if not resolutive:
        logger.info(
            '{"event": "PREPROCESSOR_NO_RESOLUTIVE_ZONE", "source": "%s"}',
            source_doc,
        )
        return []

    desc_paragraphs = [p.strip() for p in descriptive.split("\n") if p.strip()]

    items: List[ViolationWithContext] = []
    for m in _NUMBERED_ITEM_RE.finditer(resolutive):
        statement = m.group(2).strip()

        lm = _LAW_REF_RE.search(statement)
        law_ref = lm.group(0).strip() if lm else ""
        subject = _extract_subject(statement)

        ctx = _find_context_in_descriptive(law_ref, desc_paragraphs, statement)

        items.append(ViolationWithContext(
            violation_id=str(uuid.uuid4()),
            statement=statement,
            table_formulation=statement,
            context=ctx,
            section="описательная",
            page=1,
            source_document=source_doc,
            law_ref=law_ref,
            subject=subject,
        ))

    return items


# ---------------------------------------------------------------------------
# LLM fallback
# ---------------------------------------------------------------------------


def _llm_fallback(raw_text: str, source_doc: str) -> List[ViolationWithContext]:
    """LLM fallback: use ACT_EXTRACTION_TEMPLATE when structural extraction fails."""
    from prompts import ACT_EXTRACTION_TEMPLATE, parse_extraction_blocks
    from rag_pipeline import get_pipeline

    pipeline = get_pipeline()
    chunk_size = 3000
    chunks = [raw_text[i:i + chunk_size] for i in range(0, len(raw_text), chunk_size)]

    items: List[ViolationWithContext] = []
    for chunk in chunks:
        prompt = ACT_EXTRACTION_TEMPLATE.format(act_text=chunk)
        try:
            response = pipeline.analyze(prompt)
        except Exception as exc:
            logger.error(
                '{"event": "PREPROCESSOR_LLM_CHUNK_ERROR", "error": "%s"}',
                str(exc),
            )
            continue

        blocks = parse_extraction_blocks(response)
        for block in blocks:
            text = block.get("violation", "")
            if not text:
                continue
            # In LLM mode there is no separate context — use the violation text for both
            items.append(ViolationWithContext(
                violation_id=str(uuid.uuid4()),
                statement=text,
                table_formulation=text,
                context=text,
                section=block.get("section", "описательная"),
                page=1,
                source_document=source_doc,
                law_ref=block.get("law_ref", ""),
                subject=block.get("subject", ""),
            ))

    return items


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def preprocess_act(
    raw_text: str,
    source_doc: str,
    file_path: Optional[str] = None,
    use_llm_fallback: bool = True,
) -> List[ViolationWithContext]:
    """Main entry point for violation extraction.

    Three-path strategy (stops at first success):
    1. DOCX table extraction — reads violations from the structured table directly
    2. Text segmentation — regex on the resolutive zone only
    3. LLM fallback — full document passed to the LLM

    Args:
        raw_text: Document text (from loader.load_document).
        source_doc: Document filename for attribution.
        file_path: Optional path to original file.  If it ends with .docx,
                   enables the structured table extraction path.
        use_llm_fallback: Whether to call the LLM when structural paths fail.

    Returns:
        List of ViolationWithContext, one per violation found.
    """
    # Path 1: DOCX table extraction
    if file_path and file_path.lower().endswith(".docx"):
        items = _extract_from_docx(file_path, source_doc)
        if items:
            logger.info(
                '{"event": "PREPROCESSOR_DONE", "count": %d, "method": "docx_table", "source": "%s"}',
                len(items), source_doc,
            )
            return items

    # Path 2: Text segmentation
    items = _extract_from_text(raw_text, source_doc)
    if items:
        logger.info(
            '{"event": "PREPROCESSOR_DONE", "count": %d, "method": "text_segment", "source": "%s"}',
            len(items), source_doc,
        )
        return items

    # Path 3: LLM fallback
    if use_llm_fallback:
        logger.warning(
            '{"event": "PREPROCESSOR_LLM_FALLBACK", "source": "%s"}',
            source_doc,
        )
        items = _llm_fallback(raw_text, source_doc)
        logger.info(
            '{"event": "PREPROCESSOR_DONE", "count": %d, "method": "llm_fallback", "source": "%s"}',
            len(items), source_doc,
        )
        return items

    logger.warning(
        '{"event": "PREPROCESSOR_NO_VIOLATIONS", "source": "%s"}',
        source_doc,
    )
    return []


def violations_with_context_to_violations(
    items: List[ViolationWithContext],
) -> List[Violation]:
    """Convert ViolationWithContext list to Violation list for the pipeline.

    Mapping:
        raw_text    = table_formulation  (brief formulation — used for display)
        description = context            (full descriptive block — used by agents)
    """
    result: List[Violation] = []
    for item in items:
        result.append(Violation(
            raw_text=item.table_formulation,
            source_document=item.source_document,
            page=item.page,
            section=item.section,
            description=item.context,
            subject=item.subject,
            law_ref=item.law_ref,
            id=item.violation_id,
        ))
    return result
